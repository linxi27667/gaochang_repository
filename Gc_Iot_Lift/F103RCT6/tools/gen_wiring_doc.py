# -*- coding: utf-8 -*-
"""生成 F103RCT6 类PLC控制板 接线文档（与 F407 版结构一致）"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 全局中文字体
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_heading_font(h, size):
    for run in h.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for j, htext in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = htext
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[i + 1].cells[j].text = val
    doc.add_paragraph()
    return t

# ============ 封面标题 ============
title = doc.add_heading('五款举升机板端子接线指导书（F103RCT6 类PLC控制板）', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('大剪 / 小剪 / 双柱 / 超薄 / 丝杆')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('板卡：STM32F103RCT6（类PLC控制板_2）。供电 VDD24V（DC 24V），板载 DCDC 降压至 5V/3.3V。')
doc.add_paragraph('输入：10 路光耦隔离输入（IN_SW0–IN_SW9，24V 侧），外部开关闭合导通触发，MCU 侧统一低有效，软件归一化为 1=触发。')
doc.add_paragraph('输出：6 路继电器干接点（RELAY0–RELAY5，无极性，闭合=ON）+ 4 路 24V MOSFET 高侧输出（24V_OUT0–3，负载接于 24V_OUTx 与 0V 之间）。')
doc.add_paragraph('通信：4G DTU 接 USART1（PA9/PA10，9600，无 RS485 方向脚）。存储：FRAM（I2C1，PB6/PB7）。运行灯：LED_RUN（PC13）。')

# ============ 大剪 ============
h = doc.add_heading('大剪', 1); set_heading_font(h, 16)
h = doc.add_heading('输入端子', 2); set_heading_font(h, 14)
add_table(
    ['板上端子', '功能', 'MCU 引脚', '状态'],
    [
        ['IN_SW0 / IN0', '上升按钮', 'PA1', '闭合（0V）按下触发'],
        ['IN_SW1 / IN1', '下降按钮', 'PA2', '闭合（0V）按下触发'],
        ['IN_SW2 / IN2', '锁定按钮', 'PA3', '闭合（0V）按下触发'],
        ['IN_SW3 / IN3', '急停', 'PA4', '闭合（0V）急停触发'],
        ['IN_SW4 / IN4', '主机上限位', 'PA5', '闭合到位触发'],
        ['IN_SW5 / IN5', '主机下限位', 'PA6', '闭合到位触发'],
        ['IN_SW6 / IN6', '补油按钮', 'PA7', '闭合（0V）按下触发'],
        ['IN_SW7 / IN7', '光电', 'PC4', '闭合遮挡/报警触发'],
        ['IN_SW8 / IN8', '旋转开关', 'PC5', '闭合=主机 / 断开=子机'],
        ['IN_SW9 / IN9', '子机上限位', 'PB0', '闭合到位触发'],
    ])
h = doc.add_heading('输出端子', 2); set_heading_font(h, 14)
add_table(
    ['板上端子', '功能', 'MCU 引脚', '备注'],
    [
        ['RELAY0_OUT', 'Y0 电机继电器', 'PB12（RELAY0）', '干接点闭合=ON'],
        ['24V_OUT0', 'Y1 下降阀', 'PC3（OUT0）', '24V MOSFET 高侧'],
        ['24V_OUT1', 'Y2 主机气阀', 'PC2（OUT1）', '24V MOSFET 高侧'],
        ['24V_OUT2', 'Y3 主机工作阀', 'PC1（OUT2）', '24V MOSFET 高侧'],
        ['24V_OUT3', 'Y4 子机气阀', 'PC0（OUT3）', '24V MOSFET 高侧'],
        ['RELAY1_OUT', 'Y5 子机工作阀', 'PB13（RELAY1）', '干接点（OUT 仅 4 路，第 5 阀走继电器）'],
    ])

# ============ 小剪 ============
h = doc.add_heading('小剪', 1); set_heading_font(h, 16)
h = doc.add_heading('输入端子', 2); set_heading_font(h, 14)
add_table(
    ['板上端子', '功能', 'MCU 引脚', '状态'],
    [
        ['IN_SW0 / IN0', '上升按钮', 'PA1', '闭合（0V）按下触发'],
        ['IN_SW1 / IN1', '下降按钮', 'PA2', '闭合（0V）按下触发'],
        ['IN_SW2 / IN2', '锁定按钮', 'PA3', '闭合（0V）按下触发'],
        ['IN_SW3 / IN3', '急停', 'PA4', '闭合（0V）急停触发'],
        ['IN_SW4 / IN4', '补油按钮', 'PA5', '闭合（0V）按下触发'],
        ['IN_SW5 / IN5', '上限位', 'PA6', '闭合到位触发'],
        ['IN_SW6 / IN6', '下限位', 'PA7', '闭合到位触发'],
        ['IN_SW7 / IN7', '光电', 'PC4', '闭合遮挡/报警触发'],
        ['未使用', 'IN8、IN9', 'PC5 / PB0', '不接'],
    ])
h = doc.add_heading('输出端子', 2); set_heading_font(h, 14)
add_table(
    ['板上端子', '功能', 'MCU 引脚', '备注'],
    [
        ['RELAY0_OUT', '电机继电器', 'PB12（RELAY0）', '干接点闭合=ON'],
        ['24V_OUT0', '下降阀', 'PC3（OUT0）', '24V MOSFET 高侧'],
        ['24V_OUT1', '气阀', 'PC2（OUT1）', '24V MOSFET 高侧'],
        ['未使用', 'RELAY1–5 / OUT2–3', '—', '不接'],
    ])

# ============ 双柱 ============
h = doc.add_heading('双柱', 1); set_heading_font(h, 16)
h = doc.add_heading('输入端子', 2); set_heading_font(h, 14)
add_table(
    ['板上端子', '功能', 'MCU 引脚', '状态'],
    [
        ['IN_SW0 / IN0', '上升按钮', 'PA1', '闭合（0V）按下触发'],
        ['IN_SW1 / IN1', '下降按钮', 'PA2', '闭合（0V）按下触发'],
        ['IN_SW2 / IN2', '锁定按钮', 'PA3', '闭合（0V）按下触发'],
        ['IN_SW3 / IN3', '急停', 'PA4', '闭合（0V）急停触发'],
        ['IN_SW4 / IN4', '上限位', 'PA5', '闭合到位触发'],
        ['未使用', 'IN5–IN9', 'PA6 / PA7 / PC4 / PC5 / PB0', '不接'],
    ])
h = doc.add_heading('输出端子', 2); set_heading_font(h, 14)
add_table(
    ['板上端子', '功能', 'MCU 引脚', '备注'],
    [
        ['RELAY0_OUT', '电机继电器', 'PB12（RELAY0）', '干接点闭合=ON'],
        ['24V_OUT0', '电磁铁', 'PC3（OUT0）', '24V MOSFET 高侧'],
        ['24V_OUT1', '下降阀', 'PC2（OUT1）', '24V MOSFET 高侧'],
        ['未使用', 'RELAY1–5 / OUT2–3', '—', '不接'],
    ])

# ============ 超薄 ============
h = doc.add_heading('超薄', 1); set_heading_font(h, 16)
h = doc.add_heading('输入端子', 2); set_heading_font(h, 14)
add_table(
    ['板上端子', '功能', 'MCU 引脚', '状态'],
    [
        ['IN_SW0 / IN0', '上升按钮', 'PA1', '闭合（0V）按下触发'],
        ['IN_SW1 / IN1', '下降按钮', 'PA2', '闭合（0V）按下触发'],
        ['IN_SW2 / IN2', '锁定按钮', 'PA3', '闭合（0V）按下触发'],
        ['IN_SW3 / IN3', '急停', 'PA4', '闭合（0V）急停触发'],
        ['IN_SW4 / IN4', '补油按钮', 'PA5', '闭合（0V）按下触发'],
        ['IN_SW5 / IN5', '上限位', 'PA6', '闭合到位触发'],
        ['IN_SW6 / IN6', '下限位', 'PA7', '闭合到位触发'],
        ['IN_SW7 / IN7', '光电', 'PC4', '闭合遮挡/报警触发'],
        ['未使用', 'IN8、IN9', 'PC5 / PB0', '不接'],
    ])
h = doc.add_heading('输出端子', 2); set_heading_font(h, 14)
add_table(
    ['板上端子', '功能', 'MCU 引脚', '备注'],
    [
        ['RELAY0_OUT', '电机继电器', 'PB12（RELAY0）', '干接点闭合=ON'],
        ['24V_OUT0', '下降阀', 'PC3（OUT0）', '24V MOSFET 高侧'],
        ['24V_OUT1', '气阀', 'PC2（OUT1）', '24V MOSFET 高侧'],
        ['未使用', 'RELAY1–5 / OUT2–3', '—', '不接'],
    ])

# ============ 丝杆 ============
h = doc.add_heading('丝杆（F103 新增，F407 无专表）', 1); set_heading_font(h, 16)
h = doc.add_heading('输入端子', 2); set_heading_font(h, 14)
add_table(
    ['板上端子', '功能', 'MCU 引脚', '状态'],
    [
        ['IN_SW0 / IN0', '上升键', 'PA1', '闭合（0V）按下触发'],
        ['IN_SW1 / IN1', '下降键', 'PA2', '闭合（0V）按下触发'],
        ['IN_SW2 / IN2', '停止键（预留）', 'PA3', '宏保留，业务不读，可不接'],
        ['IN_SW3 / IN3', '左上防撞', 'PA4', '闭合触发，EXTI 下降沿即停'],
        ['IN_SW4 / IN4', '右上防撞', 'PA5', '闭合触发，EXTI 下降沿即停'],
        ['IN_SW5 / IN5', '左下防撞', 'PA6', '闭合触发，EXTI 下降沿即停'],
        ['IN_SW6 / IN6', '右下防撞', 'PA7', '闭合触发，EXTI 下降沿即停'],
        ['未使用', 'IN7–IN9', 'PC4 / PC5 / PB0', '不接'],
    ])
h = doc.add_heading('输出端子', 2); set_heading_font(h, 14)
add_table(
    ['板上端子', '功能', 'MCU 引脚', '备注'],
    [
        ['RELAY0_OUT', '左电机电源', 'PB12（RELAY0）', '干接点闭合=ON'],
        ['RELAY1_OUT', '右电机电源', 'PB13（RELAY1）', '干接点闭合=ON'],
        ['RELAY2_OUT', '上升方向', 'PB14（RELAY2）', '干接点闭合=ON'],
        ['RELAY3_OUT', '下降方向', 'PB15（RELAY3）', '干接点闭合=ON'],
        ['未使用', 'RELAY4–5 / OUT0–3', '—', '不接'],
    ])
h = doc.add_heading('编码器输入', 2); set_heading_font(h, 14)
add_table(
    ['信号', 'MCU 引脚', '说明'],
    [
        ['左列编码器', 'PB8（TIM4_CH3）', '下降沿计数'],
        ['右列编码器', 'PB9（TIM4_CH4）', '下降沿计数'],
    ])

doc.save(r'E:\MCU\gaochang\Gc_Iot_Lift\F103RCT6\接线文档.docx')
print('OK')
