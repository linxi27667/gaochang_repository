#!/usr/bin/env python3
"""生成高昌机电多功能控制板硬件设计计划书 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全局样式 ──────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = '微软雅黑'
    h.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 封面 ──────────────────────────────────────────────────
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('高昌机电')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('多功能控制板')
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('硬件设计计划书')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()
doc.add_paragraph()

info_lines = [
    ('文档编号', 'GC-HW-SPEC-2026-001'),
    ('版    本', 'V1.0'),
    ('编制日期', '2026年06月09日'),
    ('编制单位', '高昌机电'),
    ('密    级', '商业机密'),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ── 文档控制表格 ──────────────────────────────────────────
doc.add_heading('文档控制', level=1)

t = doc.add_table(rows=3, cols=4, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['版本', '日期', '修改人', '修改说明']
for i, h in enumerate(headers):
    cell = t.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
t.rows[1].cells[0].text = 'V1.0'
t.rows[1].cells[1].text = '2026-06-09'
t.rows[1].cells[2].text = ''
t.rows[1].cells[3].text = '初始版本'
t.rows[2].cells[0].text = ''
t.rows[2].cells[1].text = ''
t.rows[2].cells[2].text = ''
t.rows[2].cells[3].text = ''

doc.add_paragraph()
doc.add_page_break()

# ── 1. 项目概述 ──────────────────────────────────────────
doc.add_heading('1. 项目概述', level=1)

doc.add_heading('1.1 项目背景', level=2)
doc.add_paragraph(
    '高昌机电为满足举升机设备的智能化控制需求，拟开发一款多功能控制板。'
    '该控制板将集成多路数字输入、继电器输出、串行通信、PWM比例阀控制等功能，'
    '同时支持物联网远程监控与管理。本计划书旨在向硬件设计公司完整阐述设计需求，'
    '供对方硬件工程师评估方案可行性并提出反馈意见。'
)

doc.add_heading('1.2 设计目标', level=2)
goals = [
    '以STM32F407ZET6为核心控制器，满足多外设驱动与实时控制需求',
    '具备掉电数据保护能力，保障举升数据不丢失',
    '支持12路24V数字输入、6路及以上继电器输出',
    '提供RS232屏幕接口、双路RS485通信（4G模块+预留）',
    '预留2路PWM用于比例阀控制',
    '支持外部Flash存储（可选OTA升级功能）',
    '预留WiFi模块安装位置，为未来无线通信扩展做准备',
    '强弱电严格隔离，确保EMC与安全合规',
]
for g in goals:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('1.3 目标应用场景', level=2)
doc.add_paragraph(
    '举升机设备智能控制，包括但不限于：举升动作控制、实时状态监控、'
    '远程数据上报、掉电安全保护等。'
)

doc.add_page_break()

# ── 2. 核心需求规格 ──────────────────────────────────────
doc.add_heading('2. 核心需求规格', level=1)

doc.add_heading('2.1 主控芯片', level=2)
t = doc.add_table(rows=8, cols=2, style='Table Grid')
mcu_rows = [
    ('芯片型号', 'STM32F407ZET6'),
    ('内核', 'ARM Cortex-M4，最高168MHz'),
    ('Flash', '512KB片内'),
    ('SRAM', '192KB'),
    ('封装', 'LQFP-144'),
    ('关键外设', '3xUSART、2xSPI、3xI2C、多路定时器PWM、ADC等'),
    ('工作电压', '1.8V ~ 3.6V'),
    ('设计要求', '外部8MHz晶振，32.768kHz RTC晶振；SWD调试接口预留'),
]
for i, (k, v) in enumerate(mcu_rows):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    for p in t.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True

doc.add_paragraph()

doc.add_heading('2.2 外部存储器选型', level=2)
doc.add_paragraph(
    '外部存储器用于保存举升机运行数据、参数配置等。当前有两种候选方案：'
)

t = doc.add_table(rows=3, cols=5, style='Table Grid')
headers = ['型号', '类型', '容量', '通信接口', '适用场景']
for i, h in enumerate(headers):
    t.rows[0].cells[i].text = h
    for p in t.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True

eeprom = ['AT24Cxx', 'EEPROM', '256B ~ 512KB', 'I2C', '小容量参数存储，无法支持OTA']
flash = ['W25Qxx', 'SPI NOR Flash', '4MB ~ 16MB', 'SPI', '大容量存储，支持OTA固件升级']
for i, row_data in enumerate([eeprom, flash], 1):
    for j, val in enumerate(row_data):
        t.rows[i].cells[j].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('【待确认事项】')
run.font.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run(
    ' 贵方未来是否需要OTA远程固件升级功能？如需要，外部Flash必须选用W25Qxx系列'
    '（容量不低于8MB，以存储双份固件）。如仅需数据记录功能，AT24Cxx即可满足。'
    '建议：考虑到未来功能扩展可能性，推荐采用W25Q128（16MB SPI NOR Flash）。'
)

doc.add_paragraph()

doc.add_heading('2.3 掉电检测与续流电路', level=2)
doc.add_paragraph(
    '举升机在运行过程中若发生突然断电，必须保证MCU有足够时间将关键数据（举升位置、'
    '运行状态、报警信息等）写入非易失存储器。为此需要设计专用的掉电检测与续流电路。'
)

doc.add_heading('设计要求：', level=3)
reqs = [
    '续流时间：≥ 300ms，推荐设计目标500ms（确保数据写入完成）',
    '续流范围：仅给MCU及必要外围电路（Flash、RTC）供电，不给继电器、'
    '24V输出口、屏幕等大功率电路续流，以减小储能元件容量',
    '掉电检测：通过比较器或ADC监测主电源电压，电压跌落至阈值时触发中断，'
    'MCU立即执行数据保存流程',
    '储能方案：建议采用超级电容或小型电解电容，需硬件工程师核算具体容量',
    '隔离设计：续流电路供电路径与主电路供电路径独立，通过MOS管或二极管切换',
]
for r in reqs:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('【设计建议】')
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
p.add_run(
    ' 参考实现：24V主电源→LDO→3.3V为MCU供电；掉电检测比较器接在24V侧，'
    '当24V跌落至约20V时触发MCU中断（EXTI）；超级电容储能于3.3V侧，'
    '仅覆盖MCU+Flash+RTC。请硬件工程师核算：在MCU最大功耗约100mA条件下，'
    '维持500ms所需电容容量（约50mF@3.3V）。'
)

doc.add_paragraph()

doc.add_heading('2.4 数字输入接口（24V DI）', level=2)
t = doc.add_table(rows=6, cols=2, style='Table Grid')
di_rows = [
    ('数量', '12路（可扩展预留至16路）'),
    ('输入电压', '24V DC（兼容12V~30V宽范围）'),
    ('接口形式', '端子排（建议5.08mm间距）'),
    ('输入保护', '光耦隔离 + RC滤波 + TVS管防浪涌'),
    ('逻辑电平', '经光耦隔离后转为3.3V逻辑电平送入MCU GPIO'),
    ('设计注意', '每路需独立限流电阻，典型值10kΩ~15kΩ；建议每4路共享一组光耦电源'),
]
for i, (k, v) in enumerate(di_rows):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    for p in t.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True

doc.add_paragraph()

doc.add_heading('2.5 继电器输出', level=2)
t = doc.add_table(rows=7, cols=2, style='Table Grid')
relay_rows = [
    ('数量', '至少6路（建议预留至8路）'),
    ('输出类型', '继电器干接点输出（常开+常闭）'),
    ('触点容量', '≥ 10A/250VAC 或 10A/30VDC'),
    ('驱动方式', 'MCU GPIO → 三极管/MOS管 → 继电器线圈'),
    ('保护措施', '每路继电器线圈反并联续流二极管；触点侧按负载类型加RC吸收或压敏电阻'),
    ('指示灯', '每路继电器输出状态LED指示'),
    ('接口形式', '端子排输出'),
]
for i, (k, v) in enumerate(relay_rows):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    for p in t.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True

doc.add_paragraph()
doc.add_paragraph(
    '此外需预留24V输出口，供外部传感器或其他设备取电。建议通过保险丝+TVS保护后引出，'
    '输出能力不低于500mA。'
)

doc.add_paragraph()

doc.add_heading('2.6 串行通信接口', level=2)

doc.add_heading('2.6.1 RS232接口（屏幕通讯）', level=3)
t = doc.add_table(rows=5, cols=2, style='Table Grid')
rs232_rows = [
    ('数量', '1路'),
    ('接口形式', 'DB9插头（公头，直连屏幕）'),
    ('电平转换', 'MAX3232或兼容芯片'),
    ('波特率', '默认115200bps，可配置'),
    ('连接定义', 'TX、RX、GND三线制；DB9针脚按标准RS232定义'),
]
for i, (k, v) in enumerate(rs232_rows):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    for p in t.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True

doc.add_paragraph()

doc.add_heading('2.6.2 RS485接口（4G模块 + 预留）', level=3)
t = doc.add_table(rows=7, cols=3, style='Table Grid')
headers = ['', 'RS485-1（4G模块）', 'RS485-2（预留）']
for i, h in enumerate(headers):
    t.rows[0].cells[i].text = h
    for p in t.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
rs485_data = [
    ('用途', '连接4G DTU/模块', '未来扩展预留'),
    ('电平转换', 'MAX485或兼容芯片', 'MAX485或兼容芯片'),
    ('接口形式', 'DB9插座 或 螺丝端子（二选一，PCB空间优先考虑螺丝端子）', '螺丝端子'),
    ('接线定义', 'A+、B-、GND', 'A+、B-、GND'),
    ('终端电阻', '预留120Ω跳线电阻位置', '预留120Ω跳线电阻位置'),
    ('收发控制', '自动方向控制电路（推荐）或MCU RTS引脚控制', '同左'),
]
for i, (k, v1, v2) in enumerate(rs485_data, 1):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v1
    t.rows[i].cells[2].text = v2
    for p in t.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('【PCB布局建议】')
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
p.add_run(
    ' 如PCB空间充足，RS485-1采用DB9插座方便4G模块即插即用，RS485-2采用螺丝端子；'
    '如空间紧张，两路均采用螺丝端子。4G模块选型支持DB9或接线端子两种接口形式，'
    '最终以PCB布局确定。'
)

doc.add_paragraph()

doc.add_heading('2.7 PWM输出（比例阀控制）', level=2)
t = doc.add_table(rows=5, cols=2, style='Table Grid')
pwm_rows = [
    ('数量', '至少2路（建议预留至4路）'),
    ('输出电压', '24V PWM信号（经MOS管驱动后输出）'),
    ('PWM频率', '1kHz ~ 20kHz 可调（默认1kHz）'),
    ('占空比分辨率', '至少10位（0.1%步进）'),
    ('驱动能力', '每路输出电流 ≥ 1A（满足常见比例阀驱动需求）'),
]
for i, (k, v) in enumerate(pwm_rows):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    for p in t.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True

doc.add_paragraph()
doc.add_paragraph(
    '比例阀驱动建议采用N-MOS管（如IRLZ44N）做低侧开关，MCU PWM信号经栅极驱动电阻'
    '控制MOS管通断。每路需反向续流二极管保护。'
)

doc.add_paragraph()

doc.add_heading('2.8 WiFi通信模块预留', level=2)
doc.add_paragraph(
    'PCB上需预留WiFi模块的安装位置与接口，暂不焊接，未来根据需求扩展。'
)
t = doc.add_table(rows=5, cols=2, style='Table Grid')
wifi_rows = [
    ('模块选型', 'ESP-12S / ESP32-C3 等（待定）'),
    ('接口方式', '排针座（2.54mm间距）引出：VCC、GND、TX、RX、RST、IO0'),
    ('供电', '3.3V（由板载LDO供电，预留300mA余量）'),
    ('PCB布局', '预留模块焊盘区域（约20mm×15mm），天线区域净空'),
    ('设计注意', 'WiFi天线区域PCB禁止铺铜和走线，需硬件工程师设计匹配电路'),
]
for i, (k, v) in enumerate(wifi_rows):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    for p in t.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True

doc.add_paragraph()

doc.add_heading('2.9 电源设计', level=2)
doc.add_paragraph('控制板供电及电源分配要求如下：')

t = doc.add_table(rows=6, cols=2, style='Table Grid')
psu_rows = [
    ('输入电源', '24V DC 工业电源'),
    ('电源保护', '自恢复保险丝 + TVS管 + 防反接二极管'),
    ('3.3V供电', 'DCDC降压（24V→5V）+ LDO（5V→3.3V），为MCU及数字电路供电'),
    ('继电器供电', '24V直接供电（与MCU电源独立通路）'),
    ('电源指示', '24V、5V、3.3V各级电源LED指示'),
    ('功耗估算', 'MCU系统约200mA；每路继电器约40~80mA；总计约800mA@24V（满载）'),
]
for i, (k, v) in enumerate(psu_rows):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    for p in t.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True

doc.add_page_break()

# ── 3. PCB设计要求 ────────────────────────────────────────
doc.add_heading('3. PCB设计要求', level=1)

doc.add_heading('3.1 强弱电隔离', level=2)
doc.add_paragraph(
    '这是本项目最关键的设计要求之一。PCB必须严格区分强电区域与弱电区域。'
)

t = doc.add_table(rows=4, cols=2, style='Table Grid')
iso_rows = [
    ('强电区域', '24V电源输入、继电器输出、24V输出口、比例阀PWM输出（24V侧）'),
    ('弱电区域', 'MCU、3.3V/5V数字电路、RS232/RS485通信电路、WiFi模块'),
    ('隔离措施',
     '1) 强弱电区域PCB走线间距≥3mm（有条件时≥5mm）\n'
     '2) 光耦隔离24V数字输入信号\n'
     '3) 通信接口（RS485）采用隔离芯片或隔离模块\n'
     '4) 电源分区布局，强电电源走线不穿越弱电区域\n'
     '5) 接地策略：单点接地或分区域铜皮连接'),
    ('EMC设计',
     '1) 四层板建议：Signal-GND-Power-Signal\n'
     '2) 关键信号线包地处理\n'
     '3) 每个IC电源引脚放置100nF去耦电容\n'
     '4) 晶振走线尽量短，下方保留完整地平面'),
]
for i, (k, v) in enumerate(iso_rows):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    for p in t.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True

doc.add_paragraph()

doc.add_heading('3.2 PCB尺寸与结构', level=2)
t = doc.add_table(rows=5, cols=2, style='Table Grid')
pcb_rows = [
    ('PCB层数', '建议4层板（如成本敏感可评估2层板方案）'),
    ('板厚', '1.6mm'),
    ('尺寸', '待定（需硬件工程师根据器件布局评估后给出建议尺寸）'),
    ('安装孔', '至少4个M3安装孔，四角分布'),
    ('接插件位置', '所有对外接口（DB9、端子排）统一布置于PCB板边，便于接线'),
]
for i, (k, v) in enumerate(pcb_rows):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    for p in t.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True

doc.add_page_break()

# ── 4. 接口汇总表 ────────────────────────────────────────
doc.add_heading('4. 接口资源汇总', level=1)

t = doc.add_table(rows=11, cols=5, style='Table Grid')
headers = ['序号', '接口名称', '数量', '接口形式', '备注']
for i, h in enumerate(headers):
    t.rows[0].cells[i].text = h
    for p in t.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True

iface_data = [
    ('1', '24V数字输入（DI）', '12路', '端子排', '光耦隔离'),
    ('2', '继电器输出（DO）', '≥6路', '端子排', '干接点，含LED指示'),
    ('3', '24V输出口', '≥1路', '端子排', '供外部传感器取电'),
    ('4', 'RS232屏幕接口', '1路', 'DB9公头', 'MAX3232电平转换'),
    ('5', 'RS485-1（4G模块）', '1路', 'DB9/端子排', 'MAX485，120Ω终端预留'),
    ('6', 'RS485-2（预留）', '1路', '端子排', '120Ω终端预留'),
    ('7', 'PWM比例阀输出', '≥2路', '端子排', '24V MOS驱动'),
    ('8', 'WiFi模块座', '1路', '排针座', '暂不焊接'),
    ('9', 'SWD调试接口', '1路', '2.54mm排针', 'SWDIO/SWCLK/GND/3.3V'),
    ('10', '电源输入', '1路', '端子排', '24V DC，含保护电路'),
]
for i, row in enumerate(iface_data, 1):
    for j, val in enumerate(row):
        t.rows[i].cells[j].text = val

doc.add_page_break()

# ── 5. 方案可行性评审要点 ─────────────────────────────────
doc.add_heading('5. 方案可行性评审要点', level=1)
doc.add_paragraph(
    '请硬件设计工程师针对以下要点进行可行性评估，并在方案回复中逐项给出意见：'
)

review_items = [
    'STM32F407ZET6引脚资源是否满足全部外设需求？是否存在引脚冲突？'
    '如有冲突，建议替代引脚分配方案。',
    '掉电续流电路方案：请核算在目标续流时间（300~500ms）内所需储能元件容量，'
    '并给出推荐的电路拓扑（超级电容方案 vs 电解电容方案）。',
    '12路24V输入 + 6路继电器输出的光耦/驱动电路布局是否有空间约束？'
    '是否需要增加PCB尺寸？',
    '双路RS485接口的隔离方案评估：是否需要全隔离（电源+信号隔离）？'
    '如采用半隔离（仅信号隔离），对EMC性能的影响评估。',
    '比例阀PWM驱动电路的电流能力是否满足？是否需要增加外置MOS管驱动级？',
    'WiFi模块预留区域对PCB整体布局的影响，天线区域净空要求是否可满足？',
    '强弱电隔离间距在4层板/2层板两种方案下是否均能满足安全规范？',
    '整板功耗估算与电源方案合理性评估。',
    'PCB成本估算（4层板 vs 2层板），以及首批打样建议数量。',
    '请给出预计的设计周期（原理图 → PCB Layout → 打样 → 调试）。',
]
for i, item in enumerate(review_items, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_paragraph()

# ── 6. 交付物要求 ─────────────────────────────────────────
doc.add_heading('6. 交付物要求', level=1)

deliverables = [
    '原理图设计文件（Altium Designer / KiCad 格式均可）',
    'PCB Layout设计文件',
    'Gerber制造文件及钻孔文件',
    'BOM物料清单（含推荐供应商及参考价格）',
    '硬件设计说明文档（含关键电路设计分析、功耗计算、热分析等）',
    '首批样品（数量待商定，建议不少于5片）',
    '测试报告（含各接口功能测试、电源纹波测试、EMC预测试等）',
]
for d in deliverables:
    doc.add_paragraph(d, style='List Bullet')

doc.add_paragraph()

# ── 7. 时间节点 ───────────────────────────────────────────
doc.add_heading('7. 项目时间节点（建议）', level=1)

t = doc.add_table(rows=6, cols=4, style='Table Grid')
headers = ['阶段', '内容', '预计周期', '备注']
for i, h in enumerate(headers):
    t.rows[0].cells[i].text = h
    for p in t.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
timeline = [
    ('方案评审', '硬件工程师评估可行性，反馈意见', '1周', '含本计划书评审'),
    ('原理图设计', '完成原理图并评审', '2周', '需我方确认后进入Layout'),
    ('PCB Layout', 'PCB布局布线', '2周', '含DFM审查'),
    ('打样制造', 'PCB打样 + SMT贴装', '1~2周', '首批5片'),
    ('调试验证', '功能测试、问题修复', '2周', '含整改一次'),
]
for i, (phase, content, duration, note) in enumerate(timeline, 1):
    t.rows[i].cells[0].text = phase
    t.rows[i].cells[1].text = content
    t.rows[i].cells[2].text = duration
    t.rows[i].cells[3].text = note

doc.add_paragraph()
doc.add_paragraph('注：以上时间为预估，具体以硬件设计公司评估后确认的排期为准。')

doc.add_page_break()

# ── 8. 附录 ───────────────────────────────────────────────
doc.add_heading('8. 附录', level=1)

doc.add_heading('8.1 缩略语', level=2)
t = doc.add_table(rows=9, cols=2, style='Table Grid')
headers = ['缩略语', '含义']
for i, h in enumerate(headers):
    t.rows[0].cells[i].text = h
    for p in t.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
abbrs = [
    ('DI', 'Digital Input，数字输入'),
    ('DO', 'Digital Output，数字输出'),
    ('PWM', 'Pulse Width Modulation，脉冲宽度调制'),
    ('OTA', 'Over-The-Air，空中固件升级'),
    ('MCU', 'Microcontroller Unit，微控制器'),
    ('RS232/RS485', '串行通信标准'),
    ('EMC', 'Electromagnetic Compatibility，电磁兼容'),
    ('TVS', 'Transient Voltage Suppressor，瞬态电压抑制器'),
]
for i, (abbr, desc) in enumerate(abbrs, 1):
    t.rows[i].cells[0].text = abbr
    t.rows[i].cells[1].text = desc

doc.add_paragraph()

doc.add_heading('8.2 参考资料', level=2)
refs = [
    'STM32F407ZET6 Datasheet (STMicroelectronics)',
    'STM32F407xx Reference Manual (RM0090)',
    'W25Q128FV Datasheet (Winbond)',
    'MAX3232 / MAX485 Datasheet',
    'GB/T 17626 电磁兼容系列标准',
]
for r in refs:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()

# ── 签署栏 ────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 文档结束 —')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()
doc.add_paragraph()

t = doc.add_table(rows=3, cols=4, style='Table Grid')
sign_labels = [
    ('编制', '', '审核', ''),
    ('日期', '', '日期', ''),
    ('签字', '', '签字', ''),
]
for i, (l1, v1, l2, v2) in enumerate(sign_labels):
    t.rows[i].cells[0].text = l1
    t.rows[i].cells[1].text = v1
    t.rows[i].cells[2].text = l2
    t.rows[i].cells[3].text = v2
    for p in t.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True
    for p in t.rows[i].cells[2].paragraphs:
        for r in p.runs:
            r.font.bold = True

# ── 保存 ──────────────────────────────────────────────────
out_dir = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(out_dir, '高昌机电_多功能控制板_硬件设计计划书_V1.0.docx')
doc.save(out_path)
print(f'文档已生成: {out_path}')
